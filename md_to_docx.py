#!/usr/bin/env python3
"""
Markdown to Word Document Converter
Converts all markdown files in governance-docs folder to Word documents.
"""

import os
import re
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def setup_document_styles(doc):
    """Configure document styles for professional appearance."""
    styles = doc.styles

    # Modify Heading 1
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)

    # Modify Heading 2
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(16)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 76, 153)

    # Modify Heading 3
    h3_style = styles['Heading 3']
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(51, 102, 153)

    # Modify Heading 4
    h4_style = styles['Heading 4']
    h4_style.font.size = Pt(12)
    h4_style.font.bold = True
    h4_style.font.color.rgb = RGBColor(51, 102, 153)

    return doc


def process_inline_formatting(paragraph, text):
    """Process inline markdown formatting (bold, italic, code)."""
    # Pattern to match bold, italic, inline code
    patterns = [
        (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),  # Bold italic
        (r'\*\*(.+?)\*\*', 'bold'),              # Bold
        (r'\*(.+?)\*', 'italic'),                # Italic
        (r'`(.+?)`', 'code'),                    # Inline code
    ]

    # Simple approach: just add text with basic formatting detection
    # Split by formatting markers and process
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def add_table_from_html(doc, table_soup):
    """Convert HTML table to Word table."""
    rows = table_soup.find_all('tr')
    if not rows:
        return

    # Count columns from first row
    first_row = rows[0]
    cols = first_row.find_all(['th', 'td'])
    num_cols = len(cols)

    if num_cols == 0:
        return

    # Create table
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for col_idx, cell in enumerate(cells):
            if col_idx < num_cols:
                table_cell = table.rows[row_idx].cells[col_idx]
                cell_text = cell.get_text().strip()
                table_cell.text = cell_text

                # Bold headers
                if cell.name == 'th' or row_idx == 0:
                    for paragraph in table_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    # Add spacing after table
    doc.add_paragraph()


def convert_markdown_to_docx(md_content, output_path, title="Document"):
    """Convert markdown content to a Word document."""
    doc = Document()
    setup_document_styles(doc)

    # Convert markdown to HTML first
    html = markdown.markdown(
        md_content,
        extensions=['tables', 'fenced_code', 'toc']
    )
    soup = BeautifulSoup(html, 'html.parser')

    # Process elements
    for element in soup.children:
        if element.name is None:
            continue

        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            doc.add_heading(element.get_text(), level=4)
        elif element.name == 'h5':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.bold = True
            run.font.size = Pt(11)
        elif element.name == 'h6':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.bold = True
            run.font.size = Pt(10)
        elif element.name == 'p':
            p = doc.add_paragraph()
            process_inline_formatting(p, element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                # Check for checkbox
                text = li.get_text()
                if text.startswith('[ ]'):
                    p.add_run('☐ ' + text[3:].strip())
                elif text.startswith('[x]') or text.startswith('[X]'):
                    p.add_run('☑ ' + text[3:].strip())
                else:
                    process_inline_formatting(p, text)
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number')
                process_inline_formatting(p, li.get_text())
        elif element.name == 'table':
            add_table_from_html(doc, element)
        elif element.name == 'pre':
            # Code block
            code = element.find('code')
            if code:
                code_text = code.get_text()
            else:
                code_text = element.get_text()

            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)
        elif element.name == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(element.get_text())
            run.italic = True
        elif element.name == 'hr':
            # Add a horizontal line (approximation)
            p = doc.add_paragraph()
            p.add_run('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    doc.save(output_path)
    return output_path


def main():
    """Main function to convert all governance docs."""
    source_dir = Path('/Users/govind/AIEngineering/governance-docs')
    output_dir = Path('/Users/govind/AIEngineering/governance-docs-word')

    # Create output directory
    output_dir.mkdir(exist_ok=True)

    # Get all markdown files
    md_files = sorted(source_dir.glob('*.md'))

    print(f"Found {len(md_files)} markdown files to convert")
    print("-" * 50)

    for md_file in md_files:
        print(f"Converting: {md_file.name}")

        # Read markdown content
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Generate output filename
        output_name = md_file.stem + '.docx'
        output_path = output_dir / output_name

        try:
            convert_markdown_to_docx(md_content, str(output_path), md_file.stem)
            print(f"  ✓ Created: {output_path.name}")
        except Exception as e:
            print(f"  ✗ Error: {e}")

    print("-" * 50)
    print(f"Conversion complete! Files saved to: {output_dir}")


if __name__ == '__main__':
    main()
